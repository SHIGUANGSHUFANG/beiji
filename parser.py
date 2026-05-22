"""
parser.py
---------
解析背记系统 Word 文档，自动识别"知识段落"与"选择题"。

支持的文档格式
--------------

### 知识段落
  - Heading 3 / Heading 4 / Heading 5 视为标题层级
  - Normal 段落中，Word 下划线格式（[U]...[/U]）的词为填空词
  - 纯文本下划线 ___词语___ 也可识别为填空词

### 练习题区域
  - "📝 练习题" 标记开始
  - 【第X题】：  标记题号
  - 【题目内容】：题干
  - 【正确选项】：正确选项
  - 【错误选项1~3】：干扰选项

### 传统 A./B./C./D. 选择题格式也兼容

输出格式
--------
返回 list[dict]，每条记录类型为：
  {"type": "knowledge", "title": str, "content": str, "template": str, "blanks": list[str], "section": str}
  {"type": "question",  "stem": str, "options": dict, "answer": str, "section": str}
"""

import re
import random
import json
from pathlib import Path
from typing import Optional

try:
    from docx import Document
except ImportError:
    Document = None

# ─────────────────────────────────────────────
# DeepSeek API 配置
# ─────────────────────────────────────────────
# 请在此处填入你的 DeepSeek API Key
DEEPSEEK_API_KEY = "sk-78e837829f1b4b5f90d9d375dba2ecd4"   # ← 在这里填入你的 DeepSeek API Key
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"
DEEPSEEK_MODEL = "deepseek-chat"


# ─────────────────────────────────────────────
# 常量与正则
# ─────────────────────────────────────────────

# 背记文档中的标记行
EXERCISE_MARKER_RE = re.compile(r'📝\s*练习题')
QUESTION_TITLE_RE  = re.compile(r'【第[一二三四五六七八九十\d]+题】[：:]?')
QUESTION_CONTENT_RE = re.compile(r'【题目内容】[：:]\s*(.*)')
CORRECT_OPT_RE     = re.compile(r'【正确选项】[：:]\s*(.*)')
WRONG_OPT_RE       = re.compile(r'【错误选项(\d)】[：:]\s*(.*)')
SECTION_DIVIDER_RE = re.compile(r'^={5,}$')

# 传统 A./B./C./D. 选择题格式
OPTION_RE = re.compile(r'^([A-Da-d])[.、．\)）]\s*(.+)')
ANSWER_RE = re.compile(r'【?答案[】:]?\s*[:：]?\s*([A-Da-d])', re.IGNORECASE)

# 下划线填空
BLANK_TAG_RE = re.compile(r'<blank>(.*?)</blank>', re.IGNORECASE)


# ─────────────────────────────────────────────
# 段落提取
# ─────────────────────────────────────────────

def _get_paragraph_info(para) -> dict:
    """
    提取一个段落的完整信息：
    - text: 纯文本
    - styled_text: 带下划线标记的文本（[U]...[/U]）
    - style: Word 样式名
    - is_heading: 是否为标题
    - heading_level: 标题层级 (3/4/5, 0 表示非标题)
    """
    style_name = para.style.name if para.style else "Normal"
    is_heading = style_name.startswith("Heading")
    heading_level = 0
    if is_heading:
        try:
            heading_level = int(style_name.replace("Heading", "").strip())
        except ValueError:
            heading_level = 1

    # 构建 styled_text，标记下划线 run
    parts = []
    for run in para.runs:
        text = run.text
        if run.underline and text.strip():
            parts.append(f'[U]{text}[/U]')
        else:
            parts.append(text)

    styled_text = ''.join(parts)
    plain_text = styled_text.replace('[U]', '').replace('[/U]', '')

    return {
        'text': plain_text.strip(),
        'styled_text': styled_text.strip(),
        'style': style_name,
        'is_heading': is_heading,
        'heading_level': heading_level,
    }


def _extract_blanks_and_template(styled_text: str):
    """
    从带 [U]...[/U] 标记的文本中提取填空。

    返回：(template: str, blanks: list[str])
    """
    blanks = []
    counter = [0]

    def replace_underline(m):
        word = m.group(1)
        blanks.append(word)
        counter[0] += 1
        return f'[BLANK_{counter[0]}]'

    # 处理 [U]词语[/U] 格式
    template = re.sub(r'\[U\](.*?)\[/U\]', replace_underline, styled_text)

    # 处理 ___词语___ 格式（兼容纯文本下划线）
    def replace_triple(m):
        inner = m.group(1)
        blanks.append(inner)
        counter[0] += 1
        return f'[BLANK_{counter[0]}]'
    template = re.sub(r'___(.+?)___', replace_triple, template)

    # 处理纯下划线 ___ （无内容）
    def replace_pure(m):
        blanks.append('')
        counter[0] += 1
        return f'[BLANK_{counter[0]}]'
    template = re.sub(r'_{3,}', replace_pure, template)

    # 处理 <blank>词</blank>
    def replace_tag(m):
        word = m.group(1)
        blanks.append(word)
        counter[0] += 1
        return f'[BLANK_{counter[0]}]'
    template = BLANK_TAG_RE.sub(replace_tag, template)

    return template, blanks


# ─────────────────────────────────────────────
# 主解析函数
# ─────────────────────────────────────────────

def parse_docx(file_path_or_bytes) -> list[dict]:
    """
    解析 Word 文档，返回结构化条目列表。

    参数：
        file_path_or_bytes: str / Path / bytes-like（Streamlit UploadedFile）

    返回：
        list[dict]  —  每条记录含 type / content / blanks 等字段
    """
    if Document is None:
        raise ImportError("请先安装 python-docx：pip install python-docx")

    if isinstance(file_path_or_bytes, (str, Path)):
        doc = Document(str(file_path_or_bytes))
    else:
        import io
        raw = file_path_or_bytes.read() if hasattr(file_path_or_bytes, 'read') else file_path_or_bytes
        doc = Document(io.BytesIO(raw))

    # ── 第一遍：提取所有段落信息 ──
    paragraphs = []
    for para in doc.paragraphs:
        info = _get_paragraph_info(para)
        paragraphs.append(info)

    items = []
    current_section = ""          # 当前大节标题（Heading 3）
    current_subsection = ""       # 当前小节标题（Heading 4/5）
    in_exercise_zone = False      # 是否在练习题区域

    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        text = p['text']
        styled = p['styled_text']

        # ── 空行 → 跳过 ──
        if not text:
            i += 1
            continue

        # ── 分隔线 ===... ──
        if SECTION_DIVIDER_RE.match(text.strip()):
            in_exercise_zone = False
            i += 1
            continue

        # ── 练习题标记 ──
        if EXERCISE_MARKER_RE.search(text):
            in_exercise_zone = True
            i += 1
            continue

        # ── 【第X题】或【题目内容】标记 → 识别为题目 ──
        if QUESTION_TITLE_RE.match(text):
            question, consumed = _parse_question_new_format(paragraphs, i)
            if question:
                question['section'] = current_section
                items.append(question)
                i = consumed
                continue
            i += 1
            continue

        # 【题目内容】也可能没有【第X题】标题，直接出现
        if QUESTION_CONTENT_RE.match(text):
            in_exercise_zone = True  # 确保进入练习区模式
            question, consumed = _parse_question_from_content(paragraphs, i)
            if question:
                question['section'] = current_section
                items.append(question)
                i = consumed
                continue
            i += 1
            continue

        # ── 标题 ──
        if p['is_heading']:
            if p['heading_level'] == 3:
                current_section = text
                current_subsection = ""
            elif p['heading_level'] in (4, 5):
                current_subsection = text
            in_exercise_zone = False
            i += 1
            continue

        # ── 练习题区域 ──
        if in_exercise_zone:
            # 尝试匹配传统 A./B./C./D. 格式
            opt = _is_option_line(text)
            if opt:
                # 这可能是选项行，回退到题干
                if i > 0:
                    prev = paragraphs[i - 1]
                    if prev['text'] and not _is_option_line(prev['text']):
                        question, consumed = _parse_question_abcd(paragraphs, i - 1)
                        if question:
                            question['section'] = current_section
                            items.append(question)
                            i = consumed
                            continue

            i += 1
            continue

        # ── 知识段落（非练习题区域的 Normal 段落） ──
        # 将同一小节下的连续知识段落合并为一条
        if len(text) >= 2 and not in_exercise_zone:
            # 收集连续的知识段落
            knowledge_lines = []
            k = i
            while k < len(paragraphs):
                pk = paragraphs[k]
                tk = pk['text']
                sk = pk['styled_text']

                # 遇到空行，停止
                if not tk:
                    break
                # 遇到标题，停止
                if pk['is_heading']:
                    break
                # 遇到练习题标记，停止
                if EXERCISE_MARKER_RE.search(tk):
                    break
                # 遇到分隔线，停止
                if SECTION_DIVIDER_RE.match(tk.strip()):
                    break
                # 遇到选择题标记，停止
                if QUESTION_TITLE_RE.match(tk):
                    break
                # 遇到题目内容标记，停止
                if QUESTION_CONTENT_RE.match(tk):
                    break
                # 遇到正确选项标记，停止
                if CORRECT_OPT_RE.match(tk):
                    break

                knowledge_lines.append(sk)
                k += 1

            if knowledge_lines:
                # 将多条知识段落合并为一条记录
                # 如果同一条知识段落有多个子行，每行单独成一条
                for line_styled in knowledge_lines:
                    line_plain = line_styled.replace('[U]', '').replace('[/U]', '')
                    if len(line_plain) < 2:
                        continue

                    template, blanks = _extract_blanks_and_template(line_styled)
                    title = current_subsection or current_section
                    items.append({
                        'type': 'knowledge',
                        'title': title,
                        'content': template,
                        'template': template,
                        'blanks': blanks,
                        'section': current_section,
                    })

                i = k
                continue

        i += 1

    return items


# ─────────────────────────────────────────────
# 题目解析（从【题目内容】行开始，无【第X题】标题）
# ─────────────────────────────────────────────

def _parse_question_from_content(paragraphs: list, start: int) -> tuple:
    """
    从【题目内容】行开始解析一道选择题（没有【第X题】标题的情况）。

    返回 (question_dict, next_index)
    """
    stem = ""
    correct_opt = ""
    wrong_opts = []
    j = start

    # 先提取当前行的题干
    m_content = QUESTION_CONTENT_RE.match(paragraphs[j]['text'])
    if m_content:
        stem = m_content.group(1).strip()
        j += 1

    # 继续扫描后续行
    while j < len(paragraphs):
        p = paragraphs[j]
        text = p['text']

        if not text:
            # 空行可能是题与题之间的分隔
            # 但先检查后续行是否属于本题
            if correct_opt:
                # 已经有正确选项了，空行意味着本题结束
                break
            j += 1
            continue

        # 遇到下一题的各种标记 → 结束
        if QUESTION_TITLE_RE.match(text):
            break
        if QUESTION_CONTENT_RE.match(text):
            break
        if EXERCISE_MARKER_RE.search(text):
            break
        if SECTION_DIVIDER_RE.match(text.strip()):
            break
        if p['is_heading'] and not QUESTION_TITLE_RE.match(text):
            break

        # 匹配正确选项
        m_correct = CORRECT_OPT_RE.match(text)
        if m_correct:
            correct_opt = m_correct.group(1).strip()
            j += 1
            continue

        # 匹配错误选项
        m_wrong = WRONG_OPT_RE.match(text)
        if m_wrong:
            wrong_opts.append((int(m_wrong.group(1)), m_wrong.group(2).strip()))
            j += 1
            continue

        j += 1

    if not stem:
        return None, start + 1

    # 构建选项字典：A=正确, B/C/D=错误
    options = {}
    options['A'] = correct_opt
    wrong_opts.sort(key=lambda x: x[0])
    for idx, (_, content) in enumerate(wrong_opts):
        letter = chr(ord('B') + idx)
        options[letter] = content

    question = {
        'type': 'question',
        'stem': stem,
        'options': options,
        'answer': 'A',
    }
    return question, j


# ─────────────────────────────────────────────
# 题目解析（【正确选项】格式）
# ─────────────────────────────────────────────

def _parse_question_new_format(paragraphs: list, start: int) -> tuple:
    """
    解析【第X题】格式的选择题。

    返回 (question_dict, next_index)
    """
    stem = ""
    correct_opt = ""
    wrong_opts = []
    j = start

    while j < len(paragraphs):
        p = paragraphs[j]
        text = p['text']

        if not text:
            j += 1
            continue

        # 遇到下一题标记 → 结束
        if j > start and QUESTION_TITLE_RE.match(text):
            break
        # 遇到练习题标记 → 结束
        if EXERCISE_MARKER_RE.search(text):
            break
        # 遇到分隔线 → 结束
        if SECTION_DIVIDER_RE.match(text.strip()):
            break
        # 遇到标题 → 结束
        if p['is_heading']:
            break

        # 匹配题目内容
        m_content = QUESTION_CONTENT_RE.match(text)
        if m_content:
            stem = m_content.group(1).strip()
            j += 1
            continue

        # 匹配正确选项
        m_correct = CORRECT_OPT_RE.match(text)
        if m_correct:
            correct_opt = m_correct.group(1).strip()
            j += 1
            continue

        # 匹配错误选项
        m_wrong = WRONG_OPT_RE.match(text)
        if m_wrong:
            wrong_opts.append((int(m_wrong.group(1)), m_wrong.group(2).strip()))
            j += 1
            continue

        j += 1

    if not stem:
        return None, start + 1

    # 构建选项字典：A=正确, B/C/D=错误
    options = {}
    options['A'] = correct_opt
    wrong_opts.sort(key=lambda x: x[0])
    for idx, (_, content) in enumerate(wrong_opts):
        letter = chr(ord('B') + idx)
        options[letter] = content

    question = {
        'type': 'question',
        'stem': stem,
        'options': options,
        'answer': 'A',
    }
    return question, j


# ─────────────────────────────────────────────
# 题目解析（传统 A./B./C./D. 格式）
# ─────────────────────────────────────────────

def _is_option_line(text: str) -> Optional[tuple]:
    m = OPTION_RE.match(text.strip())
    if m:
        return m.group(1).upper(), m.group(2).strip()
    return None


def _is_answer_line(text: str) -> Optional[str]:
    m = ANSWER_RE.search(text)
    return m.group(1).upper() if m else None


def _parse_question_abcd(paragraphs: list, start: int) -> tuple:
    """
    解析传统 A./B./C./D. 格式的选择题。

    返回 (question_dict, next_index)
    """
    stem = paragraphs[start]['text']
    options = {}
    answer = ''
    j = start + 1

    while j < len(paragraphs) and j < start + 10:
        p = paragraphs[j]
        text = p['text']

        if not text:
            j += 1
            continue

        opt = _is_option_line(text)
        if opt:
            options[opt[0]] = opt[1]
            j += 1
            continue

        ans = _is_answer_line(text)
        if ans:
            answer = ans
            j += 1
            continue

        # 非选项、非答案、非空行 → 结束
        break

    if len(options) < 2:
        return None, start + 1

    question = {
        'type': 'question',
        'stem': stem,
        'options': options,
        'answer': answer,
    }
    return question, j


# ─────────────────────────────────────────────
# DeepSeek API 调用
# ─────────────────────────────────────────────

def _call_deepseek_api(prompt: str, system_msg: str = "你是一位初中学科教育专家。") -> str:
    """
    调用 DeepSeek API 生成内容。
    如果 API Key 未配置，返回空字符串。
    """
    if not DEEPSEEK_API_KEY:
        return ""

    try:
        import urllib.request
        import urllib.error

        payload = json.dumps({
            "model": DEEPSEEK_MODEL,
            "messages": [
                {"role": "system", "content": system_msg},
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.7,
            "max_tokens": 500,
        }).encode("utf-8")

        req = urllib.request.Request(
            DEEPSEEK_API_URL,
            data=payload,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            },
            method="POST",
        )

        with urllib.request.urlopen(req, timeout=15) as resp:
            result = json.loads(resp.read().decode("utf-8"))
            return result["choices"][0]["message"]["content"].strip()

    except Exception as e:
        print(f"[DeepSeek API 调用失败] {e}")
        return ""


def _generate_distractors_via_api(target_blanks: list[str], context: str, n: int) -> list[str]:
    """
    通过 DeepSeek API 生成干扰项。
    """
    blanks_str = "、".join(f"「{b}」" for b in target_blanks)
    prompt = (
        f"在以下知识语境中，填空处应填的词是：{blanks_str}。\n"
        f"知识语境：{context}\n\n"
        f"请生成 {n} 个干扰项（即与正确答案容易混淆但错误的词或短语），"
        f"要求：\n"
        f"1. 与正确答案在字数、学科领域上相近，容易造成混淆\n"
        f"2. 不能与正确答案相同或过于相似\n"
        f"3. 不能互相重复\n\n"
        f"请直接用 JSON 数组格式输出，例如：[\"干扰项1\", \"干扰项2\", \"干扰项3\"]\n"
        f"只输出数组，不要其他内容。"
    )

    response = _call_deepseek_api(prompt)
    if not response:
        return []

    # 尝试解析 JSON
    try:
        # 去除可能的 markdown 代码块标记
        cleaned = response.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r'^```\w*\n?', '', cleaned)
            cleaned = re.sub(r'\n?```$', '', cleaned)
            cleaned = cleaned.strip()

        result = json.loads(cleaned)
        if isinstance(result, list):
            # 过滤掉与正确答案相同的
            return [r for r in result if r not in target_blanks][:n]
    except (json.JSONDecodeError, TypeError):
        pass

    # 如果 JSON 解析失败，尝试按行拆分
    lines = [l.strip().strip('"').strip("'") for l in response.split("\n") if l.strip()]
    return [l for l in lines if l and l not in target_blanks][:n]


# ─────────────────────────────────────────────
# 生成干扰项（主函数）
# ─────────────────────────────────────────────

def generate_distractors(items, target_blanks: list[str], n: int = 3) -> list[str]:
    """
    为填空题生成干扰项。

    策略：
    1. 优先从同文档其他知识条目的 blanks 中随机选取（零延迟）
    2. 若文档内干扰词不足，且 DeepSeek API Key 已配置，则调用 AI 生成
    3. 若都不够，用占位干扰项填充

    参数：
        items: 解析后的全部条目列表（list[dict]）
        target_blanks: 当前填空的正确答案列表
        n: 需要的干扰项数量
    """
    # ── 防御性检查 ──
    if isinstance(items, (list, tuple)):
        pass  # 已经是可迭代类型
    elif hasattr(items, '__iter__'):
        # 可能是 dict_keys / dict_values / SessionStateProxy 等，安全转换
        try:
            items = list(items)
        except (TypeError, KeyError):
            items = []
    else:
        items = []
    if not items:
        items = []
    if not target_blanks:
        target_blanks = []
    if n <= 0:
        return []

    # ── 第一步：从文档内其他知识条目的填空词中抽取 ──
    pool = set()
    for item in items:
        if not isinstance(item, dict):
            continue
        if item.get('type') == 'knowledge':
            for b in item.get('blanks', []):
                if b and b not in target_blanks:
                    pool.add(b)

    pool_list = list(pool)
    random.shuffle(pool_list)

    # 如果文档内干扰词足够，直接返回
    if len(pool_list) >= n:
        return pool_list[:n]

    # ── 第二步：文档内不够，尝试调用 DeepSeek API ──
    result = pool_list[:]  # 先拿已有的

    need_from_api = n - len(result)

    if DEEPSEEK_API_KEY and need_from_api > 0:
        # 构建上下文：取当前知识条目的内容
        context = ""
        for item in items:
            if not isinstance(item, dict):
                continue
            if item.get('type') == 'knowledge' and item.get('blanks'):
                # 找到包含当前 target_blanks 的条目
                if any(b in item.get('blanks', []) for b in target_blanks):
                    context = item.get('content', '') or item.get('template', '')
                    break

        api_distractors = _generate_distractors_via_api(target_blanks, context, need_from_api + 2)
        for d in api_distractors:
            if d not in result and d not in target_blanks:
                result.append(d)
            if len(result) >= n:
                break

    # ── 第三步：还不够，用通用干扰项填充 ──
    if len(result) < n:
        generic_fillers = [
            "无", "不确定", "以上都不对", "无法判断",
            "相同", "不同", "增加", "减少",
            "变大", "变小", "升高", "降低",
        ]
        random.shuffle(generic_fillers)
        for g in generic_fillers:
            if g not in result and g not in target_blanks:
                result.append(g)
            if len(result) >= n:
                break

    return result[:n]


# ─────────────────────────────────────────────
# 示例数据（当没有上传文件时使用）
# ─────────────────────────────────────────────

DEMO_ITEMS = [
    {
        'type': 'knowledge',
        'title': '1. 现象',
        'content': '光线：用一条[BLANK_1]的线表示光传播的路径和方向。',
        'template': '光线：用一条[BLANK_1]的线表示光传播的路径和方向。',
        'blanks': ['带箭头'],
        'section': '一、光的直线传播的现象和应用',
    },
    {
        'type': 'knowledge',
        'title': '1. 现象',
        'content': '光线实际并不存在，引入方法叫做[BLANK_1]。',
        'template': '光线实际并不存在，引入方法叫做[BLANK_1]。',
        'blanks': ['理想模型法'],
        'section': '一、光的直线传播的现象和应用',
    },
    {
        'type': 'knowledge',
        'title': '2.小孔成像',
        'content': '原理：光在[BLANK_1]介质中沿直线传播',
        'template': '原理：光在[BLANK_1]介质中沿直线传播',
        'blanks': ['同种均匀'],
        'section': '一、光的直线传播的现象和应用',
    },
    {
        'type': 'knowledge',
        'title': '2.小孔成像',
        'content': '成像特点：成[BLANK_1]的[BLANK_2]；形状与原物相同，与小孔形状[BLANK_3]；大小与[BLANK_4]有关',
        'template': '成像特点：成[BLANK_1]的[BLANK_2]；形状与原物相同，与小孔形状[BLANK_3]；大小与[BLANK_4]有关',
        'blanks': ['倒立', '实像', '无关', '距离'],
        'section': '一、光的直线传播的现象和应用',
    },
    {
        'type': 'question',
        'stem': '物理学中引入"光线"来描述光的传播，主要采用了哪种研究问题的科学方法？',
        'options': {
            'A': '运用了构建理想化物理模型的方法',
            'B': '主要采用了控制变量与多次实验的方法',
            'C': '通过等效替代实际光束的方法来简化',
            'D': '它属于类比法，将光与水波类比而建立',
        },
        'answer': 'A',
        'section': '一、光的直线传播的现象和应用',
    },
    {
        'type': 'question',
        'stem': '物体在光照射下形成影子，这一现象直接表明',
        'options': {
            'A': '光在均匀介质中不会绕弯，不透明物体背后无光照射即形成影子',
            'B': '光照射物体时，物体表面发生漫反射，使反射光无法到达背面，从而产生影子',
            'C': '光在遇到物体时传播速度减慢',
            'D': '光在物体边缘发生小角度折射，背后的区域因此成为暗影',
        },
        'answer': 'A',
        'section': '一、光的直线传播的现象和应用',
    },
    {
        'type': 'question',
        'stem': '小孔成像实验能够在光屏上出现清晰的像，其根本原理是',
        'options': {
            'A': '物体上各点发出的光沿直线通过小孔，在光屏上形成了倒立的实像',
            'B': '小孔相当于一个凸透镜，对光线产生了会聚作用',
            'C': '光在穿过小孔时，频率改变导致像倒立',
            'D': '小孔可以过滤掉杂散光线，只剩下能成像的单色光',
        },
        'answer': 'A',
        'section': '一、光的直线传播的现象和应用',
    },
]
